import docx
from docx.enum.text import WD_COLOR_INDEX
import requests

def map_func(item):
    return item[:-1]
def translate(source, target, text):

    url = "https://text-translator2.p.rapidapi.com/translate"

    payload = {
        "source_language": source,
        "target_language": target,
        "text": text
    }
    headers = {
        "content-type": "application/x-www-form-urlencoded",
        "X-RapidAPI-Key": "YourRapidAPIkey",# here should be your rapidapi key
        "X-RapidAPI-Host": "text-translator2.p.rapidapi.com"
    }

    response = requests.post(url, data=payload, headers=headers)

    return response.json()['data']['translatedText']

def run_ammount(path):
    doc = docx.Document(path)
    i=0
    for para in doc.paragraphs:
        i+=len(para.runs)
    return i

def text_trans(path1,path2,lang1,lang2,rt,prob):
    
    error_text="couldn't translate text: internal server error"

    i=0
    doc = docx.Document(path1)
    print("Total runs:", run_ammount(path1))
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and run.text!=" " and ("—" not in run.text):
                ogtext=run.text
                newtext=translate(lang1,lang2,run.text)
                if (error_text in newtext) or r'\u' in newtext:
                    print("Error text caught")
                    run.text=ogtext
                    run.font.highlight_color=WD_COLOR_INDEX.RED
                else:
                    run.text=newtext
                print("ogtext= ",ogtext)
                print("newtext= ",newtext)
                print("actuall text= ",run.text)
                print("Proceccing run number ",i)
                print("------------------------")
                i=i+1
                prob['value'] = i
                rt.update_idletasks()
    doc.save(path2)

if __name__=="__main__":
    print(run_ammount("01_Огляд програми.docx"))
